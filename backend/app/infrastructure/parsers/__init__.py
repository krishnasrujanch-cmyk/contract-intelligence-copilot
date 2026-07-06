"""
Document parsers — unified interface for PDF, scanned PDF, and DOCX.

Strategy pattern: DocumentParser protocol → concrete implementations.
ParserFactory selects the correct parser by detected file type.

Security:
  - File type detected via python-magic (magic bytes), not extension
  - Max file size enforced before parsing begins
  - No shell commands — all parsing via Python libraries only
  - pytesseract uses subprocess internally; Tesseract binary is pinned in Docker
"""
from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Protocol

from app.core.config import settings
from app.core.logging import get_logger

logger = get_logger(__name__)


@dataclass
class ParseResult:
    """Unified output from any document parser."""
    text:            str                   # Full extracted text (PII not yet masked)
    tables:          list[dict[str, Any]]  # docling table dicts — atomic JSON
    image_descriptions: list[str]          # GPT-4o Vision descriptions of charts/figures
    page_count:      int
    has_images:      bool
    has_tables:      bool
    ocr_confidence:  float | None          # None if not scanned; 0-100 if OCR used
    file_type:       str                   # "pdf" | "scanned_pdf" | "docx"


class DocumentParser(Protocol):
    def parse(self, file_bytes: bytes, filename: str) -> ParseResult: ...


# ── PDF parser ────────────────────────────────────────────────────────────────

class PDFParser:
    """
    Parser for digital PDFs (text layer present).
    Uses PyMuPDF for fastest text extraction with layout preservation.
    Falls back to ScannedPDFParser if text layer is empty.
    """

    def parse(self, file_bytes: bytes, filename: str) -> ParseResult:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text: list[str] = []
        has_images = False

        for page in doc:
            pages_text.append(page.get_text("text"))
            if page.get_images():
                has_images = True

        full_text = "\n\n".join(pages_text)

        # If text layer is essentially empty → treat as scanned
        if len(full_text.strip()) < 100:
            logger.info("pdf_appears_scanned_routing_to_ocr", filename=filename)
            doc.close()
            return ScannedPDFParser().parse(file_bytes, filename)

        # Extract tables via docling (layout-aware)
        tables = self._extract_tables(file_bytes)

        doc.close()
        logger.info(
            "pdf_parsed",
            filename=filename,
            page_count=len(pages_text),
            char_count=len(full_text),
        )

        return ParseResult(
            text=full_text,
            tables=tables,
            image_descriptions=[],
            page_count=len(pages_text),
            has_images=has_images,
            has_tables=len(tables) > 0,
            ocr_confidence=None,
            file_type="pdf",
        )

    def _extract_tables(self, file_bytes: bytes) -> list[dict[str, Any]]:
        """Extract tables via docling — returns atomic JSON per table."""
        try:
            from docling.document_converter import DocumentConverter
            import tempfile, os

            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name

            try:
                converter = DocumentConverter()
                result = converter.convert(tmp_path)
                tables = []
                for i, table in enumerate(result.document.tables):
                    tables.append({
                        "index": i,
                        "data": table.export_to_dict(),
                        "raw_text": table.export_to_markdown(),
                        "page": getattr(table, "page_no", None),
                    })
                return tables
            finally:
                os.unlink(tmp_path)

        except Exception as exc:
            logger.warning("docling_table_extraction_failed", error=str(exc))
            return []


# ── Scanned PDF parser ────────────────────────────────────────────────────────

class ScannedPDFParser:
    """
    Parser for image-only (scanned) PDFs.
    Uses Tesseract OCR with per-page confidence scoring.
    Pages below threshold are flagged for Vision API fallback.
    """

    CONFIDENCE_THRESHOLD = settings.ocr_confidence_threshold

    def parse(self, file_bytes: bytes, filename: str) -> ParseResult:
        import fitz
        import pytesseract
        from PIL import Image

        doc         = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text: list[str]   = []
        confidences: list[float] = []

        for page_num, page in enumerate(doc):
            # Render page as high-res image
            matrix = fitz.Matrix(2.0, 2.0)  # 2x zoom → ~144 DPI
            pix    = page.get_pixmap(matrix=matrix)
            img    = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # OCR with confidence data
            data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
            confidences_page = [
                int(c) for c in data["conf"]
                if isinstance(c, (int, str)) and str(c).lstrip("-").isdigit() and int(c) >= 0
            ]
            avg_conf = (
                sum(confidences_page) / len(confidences_page)
                if confidences_page else 0.0
            )
            confidences.append(avg_conf)

            if avg_conf >= self.CONFIDENCE_THRESHOLD:
                page_text = pytesseract.image_to_string(img)
            else:
                # Low confidence — flag for Vision API
                logger.info(
                    "ocr_low_confidence_page",
                    page=page_num,
                    confidence=avg_conf,
                    filename=filename,
                )
                page_text = f"[OCR_LOW_CONFIDENCE page={page_num} conf={avg_conf:.1f}]"

            pages_text.append(page_text)

        overall_confidence = (
            sum(confidences) / len(confidences) if confidences else 0.0
        )

        doc.close()
        logger.info(
            "scanned_pdf_parsed",
            filename=filename,
            page_count=len(pages_text),
            avg_ocr_confidence=round(overall_confidence, 2),
        )

        return ParseResult(
            text="\n\n".join(pages_text),
            tables=[],
            image_descriptions=[],
            page_count=len(pages_text),
            has_images=True,
            has_tables=False,
            ocr_confidence=overall_confidence,
            file_type="scanned_pdf",
        )


# ── DOCX parser ───────────────────────────────────────────────────────────────

class DOCXParser:
    """
    Parser for Word documents (.docx).
    Preserves tables, headers, and paragraph structure.
    """

    def parse(self, file_bytes: bytes, filename: str) -> ParseResult:
        import docx

        doc        = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables: list[dict[str, Any]] = []

        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            tables.append({
                "index":    i,
                "data":     rows,
                "raw_text": "\n".join(" | ".join(row) for row in rows),
                "page":     None,
            })

        full_text = "\n\n".join(paragraphs)

        logger.info(
            "docx_parsed",
            filename=filename,
            paragraph_count=len(paragraphs),
            table_count=len(tables),
        )

        return ParseResult(
            text=full_text,
            tables=tables,
            image_descriptions=[],
            page_count=0,
            has_images=False,
            has_tables=len(tables) > 0,
            ocr_confidence=None,
            file_type="docx",
        )


# ── Parser factory ────────────────────────────────────────────────────────────

class ParserFactory:
    """
    Select the correct parser based on file magic bytes (not file extension).
    File extension is user-controlled and cannot be trusted.
    """

    @staticmethod
    def get_parser(file_bytes: bytes) -> DocumentParser:
        """
        Detect file type via magic bytes and return appropriate parser.

        Raises:
            ValueError: If file type is not supported or not trusted.
        """
        try:
            import magic
            mime = magic.from_buffer(file_bytes[:2048], mime=True)
        except ImportError:
            # Fallback: first bytes
            mime = ParserFactory._detect_by_header(file_bytes)

        logger.debug("detected_mime_type", mime=mime)

        if mime == "application/pdf":
            return PDFParser()
        elif mime in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            return DOCXParser()
        else:
            raise ValueError(
                f"Unsupported file type: {mime}. "
                f"Allowed: PDF, DOCX. Upload a different file."
            )

    @staticmethod
    def _detect_by_header(file_bytes: bytes) -> str:
        """Minimal header-based MIME detection without python-magic."""
        if file_bytes[:4] == b"%PDF":
            return "application/pdf"
        if file_bytes[:4] == b"PK\x03\x04":
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return "application/octet-stream"
